"""Build a copy-friendly .docx of the @atmina_lv intro thread.

Each tweet gets a heading, embedded image, and copy-paste-ready text
block. Source of truth: TWEETS list in render_intro_thread_pdf.py.
"""
import importlib.util
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DRAFTS = ROOT / "data" / "social" / "drafts"
OUT = ROOT / "docs" / "tweet_bank" / "2026-04-19-intro-thread-post.docx"

# Load TWEETS without executing render side-effects — slice the list literal
src = (ROOT / "scripts" / "render_intro_thread_pdf.py").read_text(encoding="utf-8")
start = src.find("TWEETS = [")
end = src.find("\n]\n", start) + 2
ns: dict = {}
exec(src[start:end], ns)
TWEETS = ns["TWEETS"]

URL_RE = re.compile(r"https?://\S+")


def x_weight(text: str) -> int:
    return len(URL_RE.sub("X" * 23, text))


doc = Document()

# Page-level defaults
for section in doc.sections:
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

# Title
h = doc.add_heading("atmina.lv · @atmina_lv intro thread", level=1)

meta = doc.add_paragraph()
meta.add_run("Datums: ").bold = True
meta.add_run("2026-04-19 · 5 tvīti · pirmais @atmina_lv post")

intro = doc.add_paragraph()
intro.add_run(
    "Katram tvītam pievienotā bilde ir jau ielikta šajā dokumentā — "
    "tekstu kopē no koda bloka. X limits standarta kontam: 280 simboli "
    "(URL-i X saskaita kā 23 simbolus neatkarīgi no garuma)."
)

for t in TWEETS:
    doc.add_paragraph()  # spacer
    hdr = doc.add_heading(f"Tvīts {t['n']}/5", level=2)

    img_path = DRAFTS / t["image"]
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img_path), width=Inches(6.2))

    fn = doc.add_paragraph()
    r = fn.add_run(f"Fails: {t['image']}")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    xc = x_weight(t["text"])
    status = "OK" if xc <= 280 else f"⚠️ +{xc - 280}"
    counts = doc.add_paragraph()
    cr = counts.add_run(f"X count: {xc} / 280 — {status}")
    cr.font.size = Pt(9)
    cr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph().add_run("Teksts (kopē no šejienes):").bold = True

    # Preserve paragraph breaks — split on blank lines, one docx paragraph each
    for chunk in t["text"].split("\n\n"):
        para = doc.add_paragraph()
        # Re-join any single newlines inside the chunk as soft line-breaks
        lines = chunk.split("\n")
        for i, line in enumerate(lines):
            run = para.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(11)
            if i < len(lines) - 1:
                run.add_break()

    # separator
    sep = doc.add_paragraph()
    sep.add_run("─" * 40).font.color.rgb = RGBColor(0xC0, 0xC0, 0xC0)

doc.save(str(OUT))
print(f"wrote {OUT}")
